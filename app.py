import streamlit as st
import os
import json
from agent import run_agent
from datetime import datetime

st.set_page_config(page_title="論文搜尋 Agent ouo")

STATS_FILE = "stats.json"
FAVORITES_FILE = "favorites.json"

def load_stats():
    if os.path.exists(STATS_FILE):
        with open(STATS_FILE, "r") as f:
            return json.load(f)
    return {"total": 0, "daily": {}}

def save_stats(stats):
    with open(STATS_FILE, "w") as f:
        json.dump(stats, f)

def add_click():
    stats = load_stats()
    stats["total"] += 1
    today = datetime.now().strftime("%Y-%m-%d")
    stats["daily"][today] = stats["daily"].get(today, 0) + 1
    save_stats(stats)
    return stats

def load_favorites():
    if os.path.exists(FAVORITES_FILE):
        with open(FAVORITES_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_favorites(favorites):
    with open(FAVORITES_FILE, "w", encoding="utf-8") as f:
        json.dump(favorites, f, ensure_ascii=False, indent=2)

st.markdown("""
<style>
textarea {
    border: 2px solid #2563EB !important;
    border-radius: 8px !important;
}
textarea:focus {
    border: 2px solid #1E3A5F !important;
    box-shadow: 0 0 0 2px rgba(37, 99, 235, 0.2) !important;
}
</style>
""", unsafe_allow_html=True)

os.environ["ANTHROPIC_API_KEY"] = st.secrets["ANTHROPIC_API_KEY"]

# ── 側欄 ──────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <style>
    [data-testid="stSidebar"] {
        display: flex;
        flex-direction: column;
    }
    [data-testid="stSidebarContent"] {
        display: flex;
        flex-direction: column;
        height: 100vh;
    }
    .stats-container {
        margin-top: auto;
        padding: 16px;
        background: #F0F4FF;
        border-radius: 10px;
        border: 1px solid #2563EB;
    }
    .stats-title {
        font-size: 14px;
        font-weight: bold;
        color: #1E3A5F;
        margin-bottom: 8px;
    }
    .stats-number {
        font-size: 28px;
        font-weight: bold;
        color: #2563EB;
    }
    .stats-label {
        font-size: 12px;
        color: #6B7280;
    }
    .stats-divider {
        border-top: 1px solid #D1D5DB;
        margin: 8px 0;
    }
    .stats-row {
        font-size: 12px;
        color: #374151;
        padding: 2px 0;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── 使用統計 ────────────────────────────────
    stats = load_stats()
    today = datetime.now().strftime("%Y-%m-%d")
    today_count = stats["daily"].get(today, 0)
    recent = sorted(stats["daily"].items(), reverse=True)[:5]
    recent_rows = "".join(
        f'<div class="stats-row">{date}　{count} 次</div>'
        for date, count in recent
    )

    st.markdown(f"""
    <div class="stats-container">
        <div class="stats-title">使用統計</div>
        <div class="stats-label">總搜尋次數</div>
        <div class="stats-number">{stats["total"]}</div>
        <div class="stats-divider"></div>
        <div class="stats-label">今日搜尋：<b>{today_count} 次</b></div>
        <div class="stats-divider"></div>
        <div class="stats-label" style="margin-bottom:4px">最近紀錄</div>
        {recent_rows}
    </div>
    """, unsafe_allow_html=True)
    # ── 收藏夾 ──────────────────────────────────
    st.markdown("## 收藏夾")
    favorites = load_favorites()

    if not favorites:
        st.caption("還沒有收藏的搜尋結果")
    else:
        for i, fav in enumerate(favorites):
            with st.expander(f"📄 {fav['title']}"):
                st.caption(f"關鍵字：{fav['keywords']}")
                st.caption(f"儲存時間：{fav['time']}")
                st.markdown(fav["result"])

                # 下載單筆收藏
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    st.download_button(
                        "下載 Markdown",
                        data=fav["result"],
                        file_name=f"{fav['title']}.md",
                        mime="text/markdown",
                        key=f"dl_md_{i}"
                    )
                with col_d2:
                    from docx import Document
                    from io import BytesIO
                    doc = Document()
                    doc.add_heading(fav["title"], 0)
                    for line in fav["result"].split("\n"):
                        doc.add_paragraph(line)
                    buf = BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button(
                        "下載 Word",
                        data=buf,
                        file_name=f"{fav['title']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_word_{i}"
                    )

                # 刪除單筆收藏
                if st.button("刪除", key=f"del_{i}"):
                    favorites.pop(i)
                    save_favorites(favorites)
                    st.rerun()

    st.markdown("---")

# ── 主頁面 ─────────────────────────────────────────
st.title("論文搜尋 Agent ouo")
st.markdown("### 搜尋條件")

col_a, col_b, col_c = st.columns(3)

with col_a:
    keyword1 = st.text_input("第一順位", placeholder="例如：人工智慧")
with col_b:
    keyword2 = st.text_input("第二順位", placeholder="例如：台獨")
with col_c:
    keyword3 = st.text_input("第三順位", placeholder="例如：社會階層")

if st.button("搜尋", type="primary"):
    if not keyword1.strip():
        st.warning("請至少填寫第一順位關鍵字")
    else:
        add_click()

        keywords = [k for k in [keyword1, keyword2, keyword3] if k.strip()]
        combined_query = " ".join(keywords)

        query = f"""
請搜尋以下主題的論文：
- 第一順位（必須符合）：{keyword1}
{"- 第二順位（優先符合）：" + keyword2 if keyword2 else ""}
{"- 第三順位（額外篩選）：" + keyword3 if keyword3 else ""}
搜尋時以第一順位為主要關鍵字，
回傳的論文必須與第一順位相關，
有符合第二、第三順位的論文優先排在前面。
"""

        status = st.empty()

        def update(msg):
            status.info(msg)

        with st.spinner(""):
            try:
                result = run_agent(query, progress_callback=update)
            except Exception as e:
                result = f"⚠️ 發生錯誤：{str(e)}"

        status.empty()

        if result:
            st.markdown(result)

            # 儲存到收藏夾
            st.markdown("---")
            save_title = st.text_input(
                "儲存這筆結果",
                placeholder="幫這次搜尋取個名稱，例如：AI態度相關論文",
                key="save_title"
            )
            if st.button("加入收藏夾"):
                if save_title.strip():
                    favorites = load_favorites()
                    favorites.append({
                        "title": save_title.strip(),
                        "keywords": combined_query,
                        "result": result,
                        "time": datetime.now().strftime("%Y-%m-%d %H:%M")
                    })
                    save_favorites(favorites)
                    st.success("已儲存到收藏夾！")
                else:
                    st.warning("請輸入名稱")

            # 下載
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    "下載 Markdown",
                    data=result,
                    file_name="論文搜尋.md",
                    mime="text/markdown"
                )
            with col2:
                from docx import Document
                from io import BytesIO
                doc = Document()
                doc.add_heading("論文搜尋結果", 0)
                for line in result.split("\n"):
                    doc.add_paragraph(line)
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    "下載 Word",
                    data=buf,
                    file_name="論文搜尋.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
